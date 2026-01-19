#!/usr/bin/env python3
"""
Generate test DOCX files for thesis template testing.
Creates test-njuthesis.docx and test-scut.docx with complex elements:
- Multi-level headings
- Tables
- Images
- Formulas (as formatted text)
- Chinese content
- GB/T 7714 references
"""

import os
from pathlib import Path
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from PIL import Image, ImageDraw, ImageFont


def set_chinese_font(run, font_name='SimSun', font_size=12):
    """Set Chinese font for a run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    # Set East Asian font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)


def add_heading_with_font(doc, text, level=1, font_name='SimHei'):
    """Add heading with proper Chinese font."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_chinese_font(run, font_name, 14 if level == 1 else 12)
    return heading


def add_paragraph_with_font(doc, text, font_name='SimSun', font_size=12, bold=False, alignment=None):
    """Add paragraph with proper Chinese font."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_chinese_font(run, font_name, font_size)
    run.bold = bold
    if alignment:
        para.alignment = alignment
    return para


def create_test_image(width, height, text, filename, bg_color=(240, 248, 255)):
    """Create a simple test image with text."""
    img = Image.new('RGB', (width, height), bg_color)
    draw = ImageDraw.Draw(img)

    # Draw border
    draw.rectangle([(5, 5), (width-6, height-6)], outline=(100, 100, 100), width=2)

    # Draw some shapes to make it look like a diagram
    if 'framework' in filename.lower() or 'structure' in filename.lower():
        # Draw boxes for flowchart-like diagram
        box_positions = [
            (width//2-60, 30, width//2+60, 70),
            (width//4-50, 100, width//4+50, 140),
            (width//2-50, 100, width//2+50, 140),
            (3*width//4-50, 100, 3*width//4+50, 140),
            (width//2-60, 170, width//2+60, 210),
        ]
        for box in box_positions:
            draw.rectangle(box, outline=(70, 130, 180), fill=(173, 216, 230), width=2)
        # Draw arrows
        draw.line([(width//2, 70), (width//2, 100)], fill=(100, 100, 100), width=2)
        draw.line([(width//4, 70), (width//4, 100)], fill=(100, 100, 100), width=2)
        draw.line([(3*width//4, 70), (3*width//4, 100)], fill=(100, 100, 100), width=2)

    elif 'chart' in filename.lower() or 'curve' in filename.lower():
        # Draw a simple line chart
        import random
        random.seed(42)
        points = [(50 + i*50, height - 50 - random.randint(30, 150)) for i in range(8)]
        for i in range(len(points)-1):
            draw.line([points[i], points[i+1]], fill=(70, 130, 180), width=3)
        for point in points:
            draw.ellipse([point[0]-4, point[1]-4, point[0]+4, point[1]+4], fill=(255, 100, 100))
        # Draw axes
        draw.line([(40, height-40), (width-20, height-40)], fill=(0, 0, 0), width=2)
        draw.line([(40, height-40), (40, 20)], fill=(0, 0, 0), width=2)

    elif 'architecture' in filename.lower() or 'system' in filename.lower():
        # Draw system architecture boxes
        layers = [
            (50, 30, width-50, 60, (144, 238, 144)),
            (50, 80, width-50, 130, (173, 216, 230)),
            (50, 150, width-50, 200, (255, 218, 185)),
        ]
        for box in layers:
            draw.rectangle(box[:4], outline=(100, 100, 100), fill=box[4], width=2)
        draw.line([(width//2, 60), (width//2, 80)], fill=(100, 100, 100), width=2)
        draw.line([(width//2, 130), (width//2, 150)], fill=(100, 100, 100), width=2)

    # Add text at bottom
    try:
        # Try to use a basic font
        font = ImageFont.load_default()
    except:
        font = None

    text_bbox = draw.textbbox((0, 0), text, font=font) if font else (0, 0, len(text)*6, 12)
    text_width = text_bbox[2] - text_bbox[0]
    text_x = (width - text_width) // 2
    draw.text((text_x, height - 25), text, fill=(50, 50, 50), font=font)

    # Save to BytesIO
    img_bytes = BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)
    return img_bytes


def create_njuthesis_docx(output_path):
    """Create test DOCX file for 南京大学学位论文模板."""
    doc = Document()

    # Title page info
    add_paragraph_with_font(doc, '南京大学', 'SimHei', 22, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_with_font(doc, '硕士学位论文', 'SimHei', 18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    add_paragraph_with_font(doc, '论文题目：基于深度学习的图像识别算法研究', 'SimSun', 14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_with_font(doc, 'Title: Research on Image Recognition Algorithms Based on Deep Learning', 'Times New Roman', 12, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    add_paragraph_with_font(doc, '作者姓名：张三', 'SimSun', 12)
    add_paragraph_with_font(doc, '学号：201800001', 'SimSun', 12)
    add_paragraph_with_font(doc, '专业：计算机科学与技术', 'SimSun', 12)
    add_paragraph_with_font(doc, '研究方向：人工智能', 'SimSun', 12)
    add_paragraph_with_font(doc, '指导教师：李教授', 'SimSun', 12)
    add_paragraph_with_font(doc, '学院：计算机科学与技术学院', 'SimSun', 12)

    doc.add_page_break()

    # Chinese Abstract
    add_heading_with_font(doc, '摘要', level=1)
    abstract_cn = """本文针对深度学习在图像识别领域的应用展开研究。随着人工智能技术的快速发展，
深度学习已成为计算机视觉领域最重要的技术手段之一。本研究提出了一种改进的卷积神经网络模型，
通过引入注意力机制和残差连接，有效提升了图像识别的准确率。

实验结果表明，本文提出的方法在ImageNet数据集上取得了95.6%的Top-5准确率，
相比基线模型提升了3.2个百分点。本研究的主要贡献包括：（1）提出了一种新的特征提取模块；
（2）设计了自适应的学习率调整策略；（3）构建了一个大规模中文图像数据集。

本文的研究成果对于推动深度学习在实际应用中的落地具有重要意义，
为后续研究提供了新的思路和方法。"""
    add_paragraph_with_font(doc, abstract_cn, 'SimSun', 12)

    add_paragraph_with_font(doc, '关键词：深度学习；图像识别；卷积神经网络；注意力机制；迁移学习', 'SimSun', 12, bold=True)

    doc.add_page_break()

    # English Abstract
    add_heading_with_font(doc, 'Abstract', level=1)
    abstract_en = """This thesis focuses on the application of deep learning in image recognition.
With the rapid development of artificial intelligence technology, deep learning has become one of
the most important technical means in the field of computer vision. This study proposes an improved
convolutional neural network model that effectively improves image recognition accuracy by introducing
attention mechanisms and residual connections.

Experimental results show that the method proposed in this paper achieves 95.6% Top-5 accuracy on the
ImageNet dataset, which is 3.2 percentage points higher than the baseline model. The main contributions
of this study include: (1) proposing a new feature extraction module; (2) designing an adaptive
learning rate adjustment strategy; (3) constructing a large-scale Chinese image dataset.

The research results of this paper are of great significance for promoting the practical application
of deep learning and provide new ideas and methods for subsequent research."""
    para = doc.add_paragraph(abstract_en)

    para = doc.add_paragraph()
    run = para.add_run('Keywords: ')
    run.bold = True
    para.add_run('Deep Learning; Image Recognition; Convolutional Neural Network; Attention Mechanism; Transfer Learning')

    doc.add_page_break()

    # Chapter 1
    add_heading_with_font(doc, '第一章 绪论', level=1)

    add_heading_with_font(doc, '1.1 研究背景', level=2)
    add_paragraph_with_font(doc, """随着互联网和移动设备的普及，图像数据呈现爆炸式增长。如何从海量图像中
自动提取有价值的信息，成为计算机视觉领域的核心问题。深度学习技术的出现，为图像识别带来了
革命性的突破。自2012年AlexNet在ImageNet竞赛中取得突破性成绩以来，深度学习已成为图像识别
领域的主流方法。""", 'SimSun', 12)

    add_heading_with_font(doc, '1.2 研究意义', level=2)
    add_paragraph_with_font(doc, """本研究具有重要的理论价值和实际应用意义：

（1）理论价值：本研究深入分析了深度学习模型在图像识别任务中的内在机理，
提出了新的模型架构和训练策略，丰富了深度学习的理论体系。

（2）实际应用：研究成果可应用于自动驾驶、医学影像分析、安防监控等领域，
具有广阔的应用前景。""", 'SimSun', 12)

    add_heading_with_font(doc, '1.3 研究框架', level=2)
    add_paragraph_with_font(doc, '本文的研究框架如图1-1所示：', 'SimSun', 12)

    # Add framework image
    img_bytes = create_test_image(500, 250, 'Figure 1-1: Research Framework', 'framework')
    doc.add_picture(img_bytes, width=Inches(5))
    add_paragraph_with_font(doc, '图1-1 研究框架图', 'SimSun', 10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # Chapter 2
    add_heading_with_font(doc, '第二章 相关工作', level=1)

    add_heading_with_font(doc, '2.1 传统图像识别方法', level=2)
    add_paragraph_with_font(doc, """在深度学习出现之前，传统的图像识别方法主要依赖于手工设计的特征提取器，
如SIFT、HOG、LBP等。这些方法在特定场景下取得了一定的效果，但存在以下局限性：
泛化能力有限、特征设计依赖专家经验、难以处理复杂场景。""", 'SimSun', 12)

    add_heading_with_font(doc, '2.2 深度学习方法', level=2)
    add_paragraph_with_font(doc, """近年来，深度学习方法在图像识别领域取得了显著进展。主要的网络架构包括：
VGGNet、ResNet、Inception、DenseNet等。表2-1对比了各种方法的性能。""", 'SimSun', 12)

    # Add comparison table
    add_paragraph_with_font(doc, '表2-1 主要深度学习模型性能对比', 'SimSun', 10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['模型', '年份', '层数', 'Top-1准确率(%)', 'Top-5准确率(%)']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    data = [
        ['AlexNet', '2012', '8', '57.1', '80.2'],
        ['VGGNet-16', '2014', '16', '71.5', '89.8'],
        ['ResNet-50', '2015', '50', '76.0', '93.0'],
        ['Inception-v3', '2015', '48', '78.8', '94.4'],
        ['本文方法', '2024', '52', '79.2', '95.6'],
    ]

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = value

    doc.add_paragraph()

    add_heading_with_font(doc, '2.3 损失函数', level=2)
    add_paragraph_with_font(doc, """在图像分类任务中，常用的损失函数为交叉熵损失函数。
对于多分类问题，交叉熵损失函数定义如下：""", 'SimSun', 12)

    # Formula (as formatted text)
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula_para.add_run('L = -∑ᵢ yᵢ log(pᵢ)')
    run.italic = True
    run.font.size = Pt(14)

    add_paragraph_with_font(doc, '其中，yᵢ 为真实标签，pᵢ 为预测概率。', 'SimSun', 12)

    doc.add_page_break()

    # Chapter 3
    add_heading_with_font(doc, '第三章 实验与分析', level=1)

    add_heading_with_font(doc, '3.1 数据集', level=2)
    add_paragraph_with_font(doc, """本实验使用以下数据集进行验证：

（1）ImageNet-1K：包含1000个类别，共128万张训练图像，5万张验证图像。

（2）CIFAR-100：包含100个类别，每类600张图像。

（3）自建中文图像数据集：包含50个类别的中文场景图像，共10万张。""", 'SimSun', 12)

    add_heading_with_font(doc, '3.2 实验设置', level=2)
    add_paragraph_with_font(doc, """实验环境配置如表3-1所示：""", 'SimSun', 12)

    # Add experiment settings table
    add_paragraph_with_font(doc, '表3-1 实验环境配置', 'SimSun', 10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    table2 = doc.add_table(rows=5, cols=2)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    settings = [
        ['配置项', '参数'],
        ['GPU', 'NVIDIA RTX 3090 × 4'],
        ['深度学习框架', 'PyTorch 1.12'],
        ['批量大小', '256'],
        ['学习率', '0.001 (cosine decay)'],
    ]

    for row_idx, row_data in enumerate(settings):
        for col_idx, value in enumerate(row_data):
            cell = table2.rows[row_idx].cells[col_idx]
            cell.text = value
            if row_idx == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

    doc.add_paragraph()

    add_heading_with_font(doc, '3.3 实验结果', level=2)
    add_paragraph_with_font(doc, '图3-1展示了训练过程中的准确率变化曲线：', 'SimSun', 12)

    # Add chart image
    img_bytes = create_test_image(500, 250, 'Figure 3-1: Accuracy Curves', 'chart')
    doc.add_picture(img_bytes, width=Inches(5))
    add_paragraph_with_font(doc, '图3-1 训练准确率曲线', 'SimSun', 10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # Chapter 4
    add_heading_with_font(doc, '第四章 总结与展望', level=1)

    add_heading_with_font(doc, '4.1 研究总结', level=2)
    add_paragraph_with_font(doc, """本文围绕深度学习在图像识别中的应用展开研究，主要完成了以下工作：

（1）系统分析了现有深度学习模型的优缺点，明确了改进方向。

（2）提出了一种融合注意力机制的卷积神经网络模型，有效提升了识别准确率。

（3）在多个公开数据集上进行了充分的实验验证，证明了方法的有效性。""", 'SimSun', 12)

    add_heading_with_font(doc, '4.2 未来展望', level=2)
    add_paragraph_with_font(doc, """未来的研究工作将从以下几个方面展开：

（1）探索更高效的模型压缩方法，实现模型在边缘设备上的部署。

（2）研究小样本学习和零样本学习，降低对标注数据的依赖。

（3）将本文方法拓展到视频理解、三维视觉等更复杂的任务中。""", 'SimSun', 12)

    doc.add_page_break()

    # References
    add_heading_with_font(doc, '参考文献', level=1)

    references = [
        '[1] KRIZHEVSKY A, SUTSKEVER I, HINTON G E. ImageNet classification with deep convolutional neural networks[J]. Communications of the ACM, 2017, 60(6): 84-90.',
        '[2] HE K, ZHANG X, REN S, et al. Deep residual learning for image recognition[C]//Proceedings of the IEEE conference on computer vision and pattern recognition. 2016: 770-778.',
        '[3] SIMONYAN K, ZISSERMAN A. Very deep convolutional networks for large-scale image recognition[J]. arXiv preprint arXiv:1409.1556, 2014.',
        '[4] SZEGEDY C, LIU W, JIA Y, et al. Going deeper with convolutions[C]//Proceedings of the IEEE conference on computer vision and pattern recognition. 2015: 1-9.',
        '[5] HUANG G, LIU Z, VAN DER MAATEN L, et al. Densely connected convolutional networks[C]//Proceedings of the IEEE conference on computer vision and pattern recognition. 2017: 4700-4708.',
        '[6] 李明, 王华, 张强. 基于深度学习的图像识别综述[J]. 计算机学报, 2020, 43(1): 1-24.',
        '[7] VASWANI A, SHAZEER N, PARMAR N, et al. Attention is all you need[J]. Advances in neural information processing systems, 2017, 30.',
        '[8] TAN M, LE Q. Efficientnet: Rethinking model scaling for convolutional neural networks[C]//International conference on machine learning. PMLR, 2019: 6105-6114.',
    ]

    for ref in references:
        add_paragraph_with_font(doc, ref, 'Times New Roman', 10)

    doc.add_page_break()

    # Acknowledgements
    add_heading_with_font(doc, '致谢', level=1)
    add_paragraph_with_font(doc, """时光荏苒，三年的硕士研究生生涯即将画上句号。回首这段求学岁月，
感慨万千，有太多的人需要感谢。

首先，我要衷心感谢我的导师李教授。李老师渊博的学识、严谨的治学态度、
高尚的人格魅力深深影响着我。在论文研究过程中，李老师给予了我悉心的指导和无私的帮助，
使我在科研能力和学术素养上都有了很大的提升。

其次，我要感谢实验室的各位师兄师姐和同学们。三年来，你们给予了我很多帮助，
让我在学术上受益匪浅，在生活中感受到了温暖。

最后，我要感谢我的父母，感谢他们多年来对我学业的支持和生活上的关怀。
正是有了他们的支持，我才能够专心学业，顺利完成学位论文。

张三
2024年5月于南京""", 'SimSun', 12)

    # Save document
    doc.save(output_path)
    print(f'Created: {output_path}')


def create_scut_docx(output_path):
    """Create test DOCX file for 华南理工大学博士学位论文模板."""
    doc = Document()

    # Title page info
    add_paragraph_with_font(doc, '华南理工大学', 'SimHei', 22, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_with_font(doc, '博士学位论文', 'SimHei', 18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    add_paragraph_with_font(doc, '论文题目：智能控制系统在工业机器人中的应用研究', 'SimSun', 14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_with_font(doc, 'Title: Research on Application of Intelligent Control Systems in Industrial Robots', 'Times New Roman', 12, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    add_paragraph_with_font(doc, '作者姓名：王五', 'SimSun', 12)
    add_paragraph_with_font(doc, '学号：2020301234', 'SimSun', 12)
    add_paragraph_with_font(doc, '学科专业：控制科学与工程', 'SimSun', 12)
    add_paragraph_with_font(doc, '研究方向：智能控制与机器人', 'SimSun', 12)
    add_paragraph_with_font(doc, '指导教师：陈教授', 'SimSun', 12)
    add_paragraph_with_font(doc, '学院：自动化科学与工程学院', 'SimSun', 12)

    doc.add_page_break()

    # Chinese Abstract (longer for doctoral thesis)
    add_heading_with_font(doc, '摘要', level=1)
    abstract_cn = """随着工业4.0和智能制造的快速发展，工业机器人在制造业中的应用日益广泛。
然而，传统的工业机器人控制系统存在适应性差、柔性不足等问题，难以满足复杂多变的生产环境需求。
本文针对这一问题，深入研究智能控制系统在工业机器人中的应用，旨在提高机器人的自主性、
适应性和协作能力。

本文的主要研究内容和创新点如下：

（1）提出了一种基于深度强化学习的机器人运动控制方法。该方法采用Actor-Critic架构，
通过与环境的交互学习最优控制策略，有效解决了传统控制方法在复杂任务中的局限性问题。
实验结果表明，该方法在机械臂轨迹跟踪任务中的控制精度提升了35%。

（2）设计了一种多机器人协同控制框架。基于图神经网络建模机器人之间的通信和协作关系，
实现了多机器人的分布式协同控制。该框架在多机器人搬运任务中展现出良好的可扩展性和鲁棒性。

（3）开发了一种人机协作安全控制策略。融合力/力矩传感器和视觉感知信息，
实现了机器人在与人协作过程中的自适应速度调节和碰撞规避，确保了人机协作的安全性。

（4）构建了智能控制系统实验平台。该平台集成了上述控制方法，在实际工业机器人上进行了
验证测试。实验结果证明了所提方法的有效性和实用性。

本文的研究成果为工业机器人智能化控制提供了新的理论基础和技术方案，
对推动智能制造的发展具有重要意义。"""
    add_paragraph_with_font(doc, abstract_cn, 'SimSun', 12)

    add_paragraph_with_font(doc, '关键词：智能控制；工业机器人；深度强化学习；多机器人协同；人机协作', 'SimSun', 12, bold=True)

    doc.add_page_break()

    # English Abstract (longer for doctoral thesis)
    add_heading_with_font(doc, 'Abstract', level=1)
    abstract_en = """With the rapid development of Industry 4.0 and intelligent manufacturing, industrial robots
are increasingly applied in the manufacturing industry. However, traditional industrial robot control systems
have problems such as poor adaptability and insufficient flexibility, making it difficult to meet the needs of
complex and changing production environments. This dissertation focuses on this problem and conducts in-depth
research on the application of intelligent control systems in industrial robots, aiming to improve the
autonomy, adaptability, and collaboration capabilities of robots.

The main research contents and innovations of this dissertation are as follows:

(1) A robot motion control method based on deep reinforcement learning is proposed. This method adopts the
Actor-Critic architecture and learns the optimal control strategy through interaction with the environment,
effectively solving the limitations of traditional control methods in complex tasks. Experimental results
show that this method improves the control accuracy by 35% in robotic arm trajectory tracking tasks.

(2) A multi-robot cooperative control framework is designed. Based on graph neural networks to model
the communication and cooperation relationships between robots, distributed cooperative control of
multiple robots is achieved. This framework demonstrates good scalability and robustness in multi-robot
handling tasks.

(3) A human-robot collaboration safety control strategy is developed. By integrating force/torque sensors
and visual perception information, adaptive speed adjustment and collision avoidance of robots during
collaboration with humans are realized, ensuring the safety of human-robot collaboration.

(4) An intelligent control system experimental platform is constructed. This platform integrates the
above control methods and has been verified and tested on actual industrial robots. The experimental
results demonstrate the effectiveness and practicality of the proposed methods.

The research results of this dissertation provide new theoretical foundations and technical solutions
for intelligent control of industrial robots, which is of great significance for promoting the
development of intelligent manufacturing."""
    para = doc.add_paragraph(abstract_en)

    para = doc.add_paragraph()
    run = para.add_run('Keywords: ')
    run.bold = True
    para.add_run('Intelligent Control; Industrial Robot; Deep Reinforcement Learning; Multi-robot Cooperation; Human-robot Collaboration')

    doc.add_page_break()

    # Chapter 1
    add_heading_with_font(doc, '第一章 绪论', level=1)

    add_heading_with_font(doc, '1.1 研究背景与意义', level=2)
    add_paragraph_with_font(doc, """工业机器人是智能制造的核心装备之一，在汽车制造、电子装配、物流仓储等领域
发挥着越来越重要的作用。根据国际机器人联合会(IFR)的统计数据，2022年全球工业机器人保有量
已超过350万台，中国连续多年成为全球最大的工业机器人市场。

然而，随着制造业向柔性化、智能化方向发展，传统工业机器人面临着诸多挑战：

（1）任务多样性：产品更新换代加快，机器人需要频繁切换不同的工作任务。

（2）环境复杂性：生产环境日益复杂，存在不确定性和动态变化。

（3）人机协作需求：越来越多的场景需要人与机器人协同工作。

因此，研究智能控制系统在工业机器人中的应用，提高机器人的智能化水平，
具有重要的理论价值和实际意义。""", 'SimSun', 12)

    add_heading_with_font(doc, '1.2 国内外研究现状', level=2)
    add_paragraph_with_font(doc, """近年来，国内外学者在工业机器人智能控制领域开展了大量研究工作。
在控制理论方面，模型预测控制(MPC)、自适应控制、滑模控制等方法被广泛应用于机器人运动控制。
在人工智能方面，深度学习、强化学习等技术为机器人智能化提供了新的技术路线。

表1-1总结了近年来该领域的代表性研究工作。""", 'SimSun', 12)

    # Add research summary table
    add_paragraph_with_font(doc, '表1-1 国内外研究现状总结', 'SimSun', 10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['研究者/团队', '年份', '研究内容', '主要贡献']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    data = [
        ['Levine等', '2016', '端到端学习', '首次实现机器人视觉运动控制'],
        ['OpenAI', '2019', 'Dactyl', '灵巧手操作复杂物体'],
        ['清华大学', '2020', '协作控制', '多机器人协同作业'],
        ['本文', '2024', '智能控制', '深度强化学习+人机协作'],
    ]

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = value

    doc.add_paragraph()

    add_heading_with_font(doc, '1.3 主要研究内容', level=2)
    add_paragraph_with_font(doc, '本文的研究内容框架如图1-1所示：', 'SimSun', 12)

    # Add structure image
    img_bytes = create_test_image(550, 280, 'Figure 1-1: Dissertation Structure', 'structure')
    doc.add_picture(img_bytes, width=Inches(5.5))
    add_paragraph_with_font(doc, '图1-1 论文结构图', 'SimSun', 10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # Chapter 2
    add_heading_with_font(doc, '第二章 理论基础', level=1)

    add_heading_with_font(doc, '2.1 控制理论基础', level=2)
    add_paragraph_with_font(doc, """工业机器人控制系统的设计需要以控制理论为基础。本节介绍本文研究所涉及的
基本控制理论，包括状态空间方法、最优控制理论和自适应控制理论。""", 'SimSun', 12)

    add_heading_with_font(doc, '2.1.1 状态空间方程', level=3)
    add_paragraph_with_font(doc, '对于一般的线性时不变系统，其状态空间方程可表示为：', 'SimSun', 12)

    # State space equation
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula_para.add_run('ẋ(t) = Ax(t) + Bu(t)')
    run.italic = True
    run.font.size = Pt(14)

    formula_para2 = doc.add_paragraph()
    formula_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula_para2.add_run('y(t) = Cx(t) + Du(t)')
    run.italic = True
    run.font.size = Pt(14)

    add_paragraph_with_font(doc, '其中，x(t)为状态向量，u(t)为控制输入，y(t)为系统输出，A、B、C、D为系统矩阵。', 'SimSun', 12)

    # Symbol table
    add_paragraph_with_font(doc, '表2-1 符号说明', 'SimSun', 10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    table2 = doc.add_table(rows=6, cols=3)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    symbols = [
        ['符号', '含义', '单位'],
        ['x(t)', '状态向量', '-'],
        ['u(t)', '控制输入', '-'],
        ['A', '系统矩阵', '-'],
        ['B', '输入矩阵', '-'],
        ['q', '关节角度', 'rad'],
    ]

    for row_idx, row_data in enumerate(symbols):
        for col_idx, value in enumerate(row_data):
            cell = table2.rows[row_idx].cells[col_idx]
            cell.text = value
            if row_idx == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

    doc.add_paragraph()

    add_heading_with_font(doc, '2.2 机器学习基础', level=2)
    add_paragraph_with_font(doc, """本节介绍深度学习和强化学习的基本概念和方法，为后续章节的研究奠定基础。""", 'SimSun', 12)

    add_heading_with_font(doc, '2.2.1 深度神经网络', level=3)
    add_paragraph_with_font(doc, """深度神经网络是一种具有多个隐藏层的人工神经网络。通过非线性激活函数的
层层叠加，深度网络能够学习数据的高层次抽象特征。常用的激活函数包括ReLU、Sigmoid、Tanh等。""", 'SimSun', 12)

    add_heading_with_font(doc, '2.2.2 强化学习', level=3)
    add_paragraph_with_font(doc, """强化学习是一种通过与环境交互来学习最优策略的机器学习方法。
其目标是最大化累积回报：""", 'SimSun', 12)

    formula_para3 = doc.add_paragraph()
    formula_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula_para3.add_run('J(π) = E[∑ᵗ γᵗ r(sₜ, aₜ)]')
    run.italic = True
    run.font.size = Pt(14)

    add_paragraph_with_font(doc, '其中，π为策略，γ为折扣因子，r(sₜ, aₜ)为即时奖励。', 'SimSun', 12)

    doc.add_page_break()

    # Chapter 3
    add_heading_with_font(doc, '第三章 系统设计', level=1)

    add_heading_with_font(doc, '3.1 硬件架构', level=2)
    add_paragraph_with_font(doc, """本文设计的智能控制系统采用分层架构，包括感知层、控制层和执行层。
系统硬件主要包括：工业机械臂、力/力矩传感器、深度相机、工控机等。""", 'SimSun', 12)

    add_heading_with_font(doc, '3.2 软件架构', level=2)
    add_paragraph_with_font(doc, '系统软件架构如图3-1所示：', 'SimSun', 12)

    # Add architecture image
    img_bytes = create_test_image(550, 280, 'Figure 3-1: System Architecture', 'architecture')
    doc.add_picture(img_bytes, width=Inches(5.5))
    add_paragraph_with_font(doc, '图3-1 系统软件架构图', 'SimSun', 10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_paragraph_with_font(doc, """软件系统基于ROS (Robot Operating System)开发，主要包括以下模块：

（1）感知模块：负责处理传感器数据，包括图像处理、点云处理、力信号处理等。

（2）决策模块：基于深度强化学习算法，根据感知信息生成控制决策。

（3）运动规划模块：将决策转换为具体的轨迹规划和运动控制指令。

（4）执行模块：与机器人底层控制器通信，执行运动指令。""", 'SimSun', 12)

    doc.add_page_break()

    # Chapter 4
    add_heading_with_font(doc, '第四章 实验验证', level=1)

    add_heading_with_font(doc, '4.1 实验平台搭建', level=2)
    add_paragraph_with_font(doc, """实验平台采用KUKA LBR iiwa 7协作机器人，配备ATI六维力/力矩传感器和
Intel RealSense D435深度相机。表4-1列出了实验平台的主要参数。""", 'SimSun', 12)

    # Platform parameters table
    add_paragraph_with_font(doc, '表4-1 实验平台参数', 'SimSun', 10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    table3 = doc.add_table(rows=6, cols=2)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER

    params = [
        ['参数', '数值'],
        ['机器人型号', 'KUKA LBR iiwa 7'],
        ['负载能力', '7 kg'],
        ['重复定位精度', '±0.1 mm'],
        ['自由度', '7'],
        ['控制周期', '1 ms'],
    ]

    for row_idx, row_data in enumerate(params):
        for col_idx, value in enumerate(row_data):
            cell = table3.rows[row_idx].cells[col_idx]
            cell.text = value
            if row_idx == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

    doc.add_paragraph()

    add_heading_with_font(doc, '4.2 对比实验', level=2)
    add_paragraph_with_font(doc, """为验证本文方法的有效性，设计了以下对比实验：

（1）轨迹跟踪实验：比较本文方法与传统PID控制、MPC控制的轨迹跟踪精度。

（2）抗干扰实验：在外部干扰下测试系统的鲁棒性。

（3）人机协作实验：验证安全控制策略的有效性。

表4-2展示了轨迹跟踪实验的对比结果。""", 'SimSun', 12)

    # Experiment results table
    add_paragraph_with_font(doc, '表4-2 轨迹跟踪实验结果对比', 'SimSun', 10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    table4 = doc.add_table(rows=4, cols=4)
    table4.style = 'Table Grid'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER

    results = [
        ['方法', '平均误差(mm)', '最大误差(mm)', '收敛时间(s)'],
        ['PID控制', '2.34', '5.67', '1.23'],
        ['MPC控制', '1.52', '3.21', '0.89'],
        ['本文方法', '0.98', '2.15', '0.65'],
    ]

    for row_idx, row_data in enumerate(results):
        for col_idx, value in enumerate(row_data):
            cell = table4.rows[row_idx].cells[col_idx]
            cell.text = value
            if row_idx == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

    doc.add_paragraph()

    add_paragraph_with_font(doc, '图4-1展示了轨迹跟踪实验的结果曲线：', 'SimSun', 12)

    # Add results chart
    img_bytes = create_test_image(550, 280, 'Figure 4-1: Tracking Results', 'chart_results')
    doc.add_picture(img_bytes, width=Inches(5.5))
    add_paragraph_with_font(doc, '图4-1 轨迹跟踪实验结果', 'SimSun', 10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # Chapter 5
    add_heading_with_font(doc, '第五章 结论与展望', level=1)

    add_heading_with_font(doc, '5.1 研究结论', level=2)
    add_paragraph_with_font(doc, """本文针对工业机器人智能控制问题进行了深入研究，取得了以下主要成果：

（1）提出了基于深度强化学习的机器人运动控制方法，在轨迹跟踪任务中将控制精度
提升了35%以上。

（2）设计了多机器人协同控制框架，实现了分布式协同控制，具有良好的可扩展性。

（3）开发了人机协作安全控制策略，确保了人机协作的安全性。

（4）构建了实验平台并进行了充分验证，证明了所提方法的有效性和实用性。""", 'SimSun', 12)

    add_heading_with_font(doc, '5.2 研究展望', level=2)
    add_paragraph_with_font(doc, """未来研究工作可从以下方向展开：

（1）进一步提高算法的实时性，满足高速运动控制需求。

（2）研究跨域迁移学习方法，提高算法在不同机器人平台上的通用性。

（3）探索大规模机器人集群的协同控制问题。

（4）结合大语言模型，实现自然语言驱动的机器人编程。""", 'SimSun', 12)

    doc.add_page_break()

    # References
    add_heading_with_font(doc, '参考文献', level=1)

    references = [
        '[1] LEVINE S, FINN C, DARRELL T, et al. End-to-end training of deep visuomotor policies[J]. The Journal of Machine Learning Research, 2016, 17(1): 1334-1373.',
        '[2] ANDRYCHOWICZ O M, BAKER B, CHOCIEJ M, et al. Learning dexterous in-hand manipulation[J]. The International Journal of Robotics Research, 2020, 39(1): 3-20.',
        '[3] LILLICRAP T P, HUNT J J, PRITZEL A, et al. Continuous control with deep reinforcement learning[J]. arXiv preprint arXiv:1509.02971, 2015.',
        '[4] SCHULMAN J, WOLSKI F, DHARIWAL P, et al. Proximal policy optimization algorithms[J]. arXiv preprint arXiv:1707.06347, 2017.',
        '[5] HAARNOJA T, ZHOU A, ABBEEL P, et al. Soft actor-critic: Off-policy maximum entropy deep reinforcement learning with a stochastic actor[C]//International conference on machine learning. PMLR, 2018: 1861-1870.',
        '[6] 刘金琨. 机器人控制系统的设计与MATLAB仿真[M]. 北京: 清华大学出版社, 2015.',
        '[7] 陈伯时. 自动控制系统[M]. 北京: 机械工业出版社, 2018.',
        '[8] SPONG M W, HUTCHINSON S, VIDYASAGAR M. Robot modeling and control[M]. John Wiley & Sons, 2020.',
        '[9] KOBER J, BAGNELL J A, PETERS J. Reinforcement learning in robotics: A survey[J]. The International Journal of Robotics Research, 2013, 32(11): 1238-1274.',
        '[10] 王耀南, 贺振东, 朱浩. 机器人智能控制技术研究进展[J]. 自动化学报, 2020, 46(11): 2255-2275.',
        '[11] SICILIANO B, KHATIB O. Robotics and the handbook of robotics[C]//2016 IEEE International Conference on Automation Science and Engineering. IEEE, 2016: 1229-1229.',
        '[12] BILLARD A, KRAGIC D. Trends and challenges in robot manipulation[J]. Science, 2019, 364(6446): eaat8414.',
        '[13] HADDADIN S, CROFT E. Physical human-robot interaction[M]//Springer handbook of robotics. Springer, 2016: 1835-1874.',
        '[14] 张毅, 李伟, 王天然. 多机器人协作系统研究综述[J]. 机器人, 2019, 41(5): 600-612.',
        '[15] SUTTON R S, BARTO A G. Reinforcement learning: An introduction[M]. MIT press, 2018.',
    ]

    for ref in references:
        add_paragraph_with_font(doc, ref, 'Times New Roman', 10)

    doc.add_page_break()

    # Research achievements
    add_heading_with_font(doc, '攻读博士学位期间取得的研究成果', level=1)

    add_paragraph_with_font(doc, '一、发表学术论文', 'SimHei', 12, bold=True)

    papers = [
        '[1] Wang W, Chen P, et al. Deep Reinforcement Learning for Industrial Robot Control: A Survey[J]. IEEE Transactions on Industrial Informatics, 2024, 20(3): 1234-1250. (SCI, IF=11.648, 第一作者)',
        '[2] Wang W, Chen P, et al. Multi-robot Cooperative Control Based on Graph Neural Networks[C]//IEEE International Conference on Robotics and Automation. IEEE, 2023: 5678-5685. (EI, 第一作者)',
        '[3] Wang W, Chen P, et al. Safe Human-robot Collaboration Control Strategy[J]. Robotics and Autonomous Systems, 2023, 168: 104512. (SCI, IF=4.722, 第一作者)',
    ]

    for paper in papers:
        add_paragraph_with_font(doc, paper, 'Times New Roman', 10)

    doc.add_paragraph()
    add_paragraph_with_font(doc, '二、授权专利', 'SimHei', 12, bold=True)

    patents = [
        '[1] 陈教授, 王五, 等. 一种基于深度强化学习的机器人运动控制方法[P]. 中国: CN202311234567.8, 2023-12-01.',
        '[2] 陈教授, 王五, 等. 一种多机器人协同控制系统及方法[P]. 中国: CN202310987654.3, 2023-10-15.',
    ]

    for patent in patents:
        add_paragraph_with_font(doc, patent, 'SimSun', 10)

    doc.add_paragraph()
    add_paragraph_with_font(doc, '三、参与科研项目', 'SimHei', 12, bold=True)

    projects = [
        '[1] 国家自然科学基金重点项目：智能制造系统中的多机器人协同控制理论与方法（项目编号：12345678），参与人。',
        '[2] 广东省重点研发计划：面向复杂场景的工业机器人智能控制技术研究（项目编号：2022B0101），参与人。',
    ]

    for project in projects:
        add_paragraph_with_font(doc, project, 'SimSun', 10)

    doc.add_page_break()

    # Acknowledgements
    add_heading_with_font(doc, '致谢', level=1)
    add_paragraph_with_font(doc, """时光荏苒，四年的博士研究生生涯即将落下帷幕。回首这段充实而难忘的求学岁月，
心中感慨万千，有太多的人需要感谢。

首先，我要向我的导师陈教授致以最诚挚的谢意。陈老师学识渊博、治学严谨、
为人师表，是我学术道路上的引路人。四年来，陈老师在论文选题、研究方法、
学术写作等方面给予了我悉心的指导和无私的帮助。陈老师的言传身教，
不仅让我在科研能力上有了很大的提升，更让我领悟到了做学问和做人的真谛。

感谢自动化科学与工程学院的各位老师，你们的辛勤教导让我在专业知识上
打下了坚实的基础。感谢实验室的各位师兄师姐和同门，四年来我们一起学习、
一起研究、一起成长，这段经历将是我人生中最宝贵的财富。

感谢国家自然科学基金委、广东省科技厅对本研究的资助，感谢华南理工大学
提供的良好科研条件和学习环境。

最后，我要感谢我的父母和家人。感谢你们多年来对我学业的支持和生活上的关怀，
正是有了你们的理解和支持，我才能够心无旁骛地专注于学业，顺利完成博士学位论文。
你们的爱是我前进的最大动力。

四年博士生涯虽然艰辛，但收获满满。未来的道路上，我将继续努力，
不负导师厚望，不负家人期待，在科研的道路上砥砺前行。

王五
2024年5月于广州""", 'SimSun', 12)

    # Save document
    doc.save(output_path)
    print(f'Created: {output_path}')


def main():
    """Main function to generate test DOCX files."""
    # Set up paths
    script_dir = Path(__file__).parent
    project_dir = script_dir.parent
    test_files_dir = project_dir / 'test-files'

    # Create test-files directory if it doesn't exist
    test_files_dir.mkdir(exist_ok=True)

    # Generate test files
    print('Generating test DOCX files...')
    print()

    njuthesis_path = test_files_dir / 'test-njuthesis.docx'
    scut_path = test_files_dir / 'test-scut.docx'

    create_njuthesis_docx(str(njuthesis_path))
    create_scut_docx(str(scut_path))

    print()
    print('Done! Test files created:')
    print(f'  - {njuthesis_path}')
    print(f'  - {scut_path}')
    print()
    print('You can test these files with:')
    print('  curl -X POST http://localhost:3000/thesis/upload \\')
    print(f'    -F "file=@{njuthesis_path}" \\')
    print('    -F "templateId=njuthesis"')


if __name__ == '__main__':
    main()
