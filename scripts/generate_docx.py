#!/usr/bin/env python3
"""
Generate formatted DOCX from thesis JSON data with table and image support.
Usage: python generate_docx.py <input.json> <output.docx> [--images-dir <dir>] [--template <template.docx>]
"""

import json
import sys
import os
import argparse
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


def set_chinese_font(run, font_name='宋体', size=12, bold=False):
    """Set Chinese font for a run"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    # Set East Asian font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name if font_name in ['宋体', '黑体', '楷体'] else 'Times New Roman')
    rPr.insert(0, rFonts)


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_chinese(doc, text, level, font_name='黑体', size=None, center=True):
    """Add a heading with Chinese font"""
    sizes = {0: 26, 1: 18, 2: 15, 3: 12}
    if size is None:
        size = sizes.get(level, 12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center and level <= 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_chinese_font(run, font_name, size, bold=True)

    # Add spacing
    p.paragraph_format.space_before = Pt(18 if level <= 1 else 12)
    p.paragraph_format.space_after = Pt(12 if level <= 1 else 6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    return p


def add_paragraph_chinese(doc, text, font_name='宋体', size=12, first_line_indent=True):
    """Add a paragraph with Chinese font"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, font_name, size)

    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 2 characters

    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def add_table_from_data(doc, table_data, style='Table Grid'):
    """Add a table from extracted table data"""
    rows = table_data.get('rows', [])
    if not rows:
        return None

    row_count = len(rows)
    col_count = max(len(row) for row in rows) if rows else 0

    if row_count == 0 or col_count == 0:
        return None

    table = doc.add_table(rows=row_count, cols=col_count)
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(str(cell_text))
                set_chinese_font(run, '宋体', 10)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add spacing after table
    doc.add_paragraph()

    return table


def add_image_from_file(doc, image_path, width_inches=5, caption=None):
    """Add an image from file"""
    if not os.path.exists(image_path):
        print(f"Warning: Image not found: {image_path}")
        return None

    try:
        doc.add_picture(image_path, width=Inches(width_inches))

        # Center the image
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add caption if provided
        if caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            set_chinese_font(run, '宋体', 10)

        return True
    except Exception as e:
        print(f"Warning: Failed to add image {image_path}: {e}")
        return None


def add_image_from_buffer(doc, image_buffer, width_inches=5, caption=None):
    """Add an image from buffer"""
    try:
        doc.add_picture(BytesIO(image_buffer), width=Inches(width_inches))

        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            set_chinese_font(run, '宋体', 10)

        return True
    except Exception as e:
        print(f"Warning: Failed to add image from buffer: {e}")
        return None


def generate_thesis_docx(data, output_path, images_dir=None, template_path=None):
    """Generate DOCX from thesis data"""

    # Use template if provided, otherwise create new document
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
        # Clear template content but keep styles
        for element in doc.element.body[:]:
            doc.element.body.remove(element)
    else:
        doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    metadata = data.get('metadata', {})
    tables = data.get('tables', [])
    images = data.get('images', [])

    # === Cover Page ===
    for _ in range(2):
        doc.add_paragraph()

    # University name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('上海交通大学')
    set_chinese_font(run, '黑体', 36, bold=True)

    doc.add_paragraph()

    # Thesis type
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('本科毕业论文')
    set_chinese_font(run, '黑体', 24, bold=True)

    for _ in range(2):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(metadata.get('title', '论文标题'))
    set_chinese_font(run, '黑体', 22, bold=True)

    # English title if present
    if metadata.get('title_en'):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(metadata.get('title_en'))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True

    for _ in range(3):
        doc.add_paragraph()

    # Metadata table
    info_items = [
        ('学生姓名', metadata.get('author_name', '')),
        ('学生学号', metadata.get('student_id', '')),
        ('专    业', metadata.get('major', '')),
        ('指导教师', metadata.get('supervisor', '')),
        ('学    院', metadata.get('school', '')),
    ]

    # Create a table for metadata
    info_table = doc.add_table(rows=len([i for i in info_items if i[1]]), cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    row_idx = 0
    for label, value in info_items:
        if value:
            row = info_table.rows[row_idx]
            # Label cell
            cell = row.cells[0]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(f'{label}：')
            set_chinese_font(run, '宋体', 14)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Value cell
            cell = row.cells[1]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(value)
            set_chinese_font(run, '宋体', 14)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            row_idx += 1

    # Set column widths
    for row in info_table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(6)

    for _ in range(2):
        doc.add_paragraph()

    # Date
    if metadata.get('date'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(metadata.get('date'))
        set_chinese_font(run, '宋体', 14)

    doc.add_page_break()

    # === Abstract ===
    if data.get('abstract'):
        add_heading_chinese(doc, '摘  要', 1, '黑体', 18)
        add_paragraph_chinese(doc, data['abstract'])

        if data.get('keywords'):
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run('关键词：')
            set_chinese_font(run, '黑体', 12, bold=True)
            run = p.add_run(data['keywords'])
            set_chinese_font(run, '宋体', 12)

        doc.add_page_break()

    # === English Abstract ===
    if data.get('abstract_en'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('ABSTRACT')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(18)
        run.font.bold = True
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)

        p = doc.add_paragraph()
        run = p.add_run(data['abstract_en'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        if data.get('keywords_en'):
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run('Keywords: ')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            run = p.add_run(data['keywords_en'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        doc.add_page_break()

    # === Table of Contents placeholder ===
    add_heading_chinese(doc, '目  录', 1, '黑体', 18)
    p = doc.add_paragraph()
    run = p.add_run('（目录将在 Word 中自动生成）')
    set_chinese_font(run, '宋体', 12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # === Sections ===
    table_index = 0
    image_index = 0

    for section in data.get('sections', []):
        level = section.get('level', 1)
        title = section.get('title', '')
        content = section.get('content', '')

        # Add heading based on level
        if level == 1:
            add_heading_chinese(doc, title, 1, '黑体', 16, center=True)
        elif level == 2:
            add_heading_chinese(doc, title, 2, '黑体', 14, center=False)
        else:
            add_heading_chinese(doc, title, 3, '黑体', 12, center=False)

        # Add content paragraphs
        if content:
            # Check for table placeholders like {%table_1%}
            import re
            parts = re.split(r'\{%table_(\d+)%\}', content)

            for i, part in enumerate(parts):
                if i % 2 == 0:
                    # Text content
                    paragraphs = part.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            add_paragraph_chinese(doc, para.strip())
                else:
                    # Table placeholder
                    tbl_num = int(part)
                    if tbl_num <= len(tables):
                        add_table_from_data(doc, tables[tbl_num - 1])

            # Check for image placeholders like {%img_1%}
            # Add images if referenced in content
            img_matches = re.findall(r'\{%(?:img|media)_(\d+)%\}', content)
            for img_num in img_matches:
                idx = int(img_num) - 1
                if idx < len(images) and images_dir:
                    img_info = images[idx]
                    img_path = os.path.join(images_dir, img_info.get('filename', ''))
                    if os.path.exists(img_path):
                        add_image_from_file(doc, img_path, caption=f"图 {img_num}")

    # === Add remaining tables if not placed ===
    # (Tables that weren't referenced in content)

    # === References ===
    if data.get('references'):
        doc.add_page_break()
        add_heading_chinese(doc, '参考文献', 1, '黑体', 18)

        refs = data['references']
        if isinstance(refs, str):
            # Split by newlines or reference numbers
            ref_lines = refs.strip().split('\n')
            for line in ref_lines:
                if line.strip():
                    add_paragraph_chinese(doc, line.strip(), first_line_indent=False)
        elif isinstance(refs, list):
            for i, ref in enumerate(refs, 1):
                add_paragraph_chinese(doc, f'[{i}] {ref}', first_line_indent=False)

    # === Acknowledgements ===
    if data.get('acknowledgements'):
        doc.add_page_break()
        add_heading_chinese(doc, '致  谢', 1, '黑体', 18)
        add_paragraph_chinese(doc, data['acknowledgements'])

    # Save
    doc.save(output_path)
    print(f'Generated: {output_path}')
    return output_path


def main():
    parser = argparse.ArgumentParser(description='Generate thesis DOCX from JSON data')
    parser.add_argument('input', help='Input JSON file path')
    parser.add_argument('output', help='Output DOCX file path')
    parser.add_argument('--images-dir', help='Directory containing extracted images')
    parser.add_argument('--template', help='Word template file for styling')

    args = parser.parse_args()

    with open(args.input, 'r', encoding='utf-8') as f:
        data = json.load(f)

    generate_thesis_docx(
        data,
        args.output,
        images_dir=args.images_dir,
        template_path=args.template
    )


if __name__ == '__main__':
    main()
